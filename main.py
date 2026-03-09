import json
from io import BytesIO
import os
import re
from dataclasses import dataclass
from typing import Any, Dict, List, Optional

import requests
from dotenv import load_dotenv
from flask import Flask, jsonify, render_template, request, send_file
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import simpleSplit
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas


load_dotenv()


DASHSCOPE_API_KEY = os.getenv("DASHSCOPE_API_KEY", "").strip()
QWEN_MODEL = os.getenv("QWEN_MODEL", "qwen-plus")
QWEN_BASE_URL = os.getenv(
    "QWEN_BASE_URL", "https://dashscope.aliyuncs.com/compatible-mode/v1"
)
WAN_MODEL = os.getenv("WAN_MODEL", "wan2.6-t2v")
WAN_BASE_URL = os.getenv("WAN_BASE_URL", "https://dashscope.aliyuncs.com/api/v1")


app = Flask(__name__, template_folder="templates", static_folder="static")


@dataclass
class AgentRequest:
    stage: str
    payload: Dict[str, Any]


def _extract_json(text: str) -> Dict[str, Any]:
    """Extract the first JSON object from model output and parse it safely."""
    text = text.strip()

    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?", "", text).strip()
        text = re.sub(r"```$", "", text).strip()

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    match = re.search(r"\{[\s\S]*\}", text)
    if not match:
        raise ValueError("Model did not return JSON.")

    return json.loads(match.group(0))


def _call_qwen_json(system_prompt: str, user_prompt: str) -> Dict[str, Any]:
    if not DASHSCOPE_API_KEY:
        raise RuntimeError("DASHSCOPE_API_KEY is missing in .env")

    url = f"{QWEN_BASE_URL}/chat/completions"
    headers = {
        "Authorization": f"Bearer {DASHSCOPE_API_KEY}",
        "Content-Type": "application/json",
    }
    body = {
        "model": QWEN_MODEL,
        "temperature": 0.7,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    }

    response = requests.post(url, headers=headers, json=body, timeout=60)
    response.raise_for_status()
    data = response.json()
    content = data["choices"][0]["message"]["content"]
    return _extract_json(content)


def _call_qwen_text(system_prompt: str, user_prompt: str) -> str:
    if not DASHSCOPE_API_KEY:
        raise RuntimeError("DASHSCOPE_API_KEY is missing in .env")

    url = f"{QWEN_BASE_URL}/chat/completions"
    headers = {
        "Authorization": f"Bearer {DASHSCOPE_API_KEY}",
        "Content-Type": "application/json",
    }
    body = {
        "model": QWEN_MODEL,
        "temperature": 0.8,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    }

    response = requests.post(url, headers=headers, json=body, timeout=60)
    response.raise_for_status()
    data = response.json()
    return data["choices"][0]["message"]["content"].strip()


def _video_script_prompt(payload: Dict[str, Any]) -> str:
    return f"""
你是短剧编导，请直接输出一段可拍可生视频的短剧脚本。

输入信息：
- 题材：{payload.get('genre', '')}
- 核心设定：{payload.get('idea', '')}
- 人物：{payload.get('roles', '')}
- 风格：{payload.get('style', '')}
- 时长：{payload.get('duration_sec', 10)} 秒

输出要求：
1) 先给出“标题”。
2) 再给“短剧脚本（分镜级）”，包含 4-6 个镜头，每个镜头写画面、动作、台词/音效。
3) 最后给“视频生成提示词（中文）”，用于文生视频，保证画面连贯。
4) 全文中文，简洁有戏剧冲突。
""".strip()


def _create_video_task(payload: Dict[str, Any]) -> Dict[str, Any]:
    if not DASHSCOPE_API_KEY:
        raise RuntimeError("DASHSCOPE_API_KEY is missing in .env")

    prompt = str(payload.get("prompt", "")).strip()
    if not prompt:
        raise ValueError("Video prompt is required.")

    model = payload.get("model", WAN_MODEL)
    size = payload.get("size", "1280*720")
    duration = int(payload.get("duration", 10))
    prompt_extend = bool(payload.get("prompt_extend", True))

    url = f"{WAN_BASE_URL}/services/aigc/video-generation/video-synthesis"
    headers = {
        "Authorization": f"Bearer {DASHSCOPE_API_KEY}",
        "Content-Type": "application/json",
        "X-DashScope-Async": "enable",
    }
    body = {
        "model": model,
        "input": {"prompt": prompt},
        "parameters": {
            "size": size,
            "duration": duration,
            "prompt_extend": prompt_extend,
            "watermark": False,
        },
    }

    response = requests.post(url, headers=headers, json=body, timeout=60)
    response.raise_for_status()
    return response.json()


def _query_video_task(task_id: str) -> Dict[str, Any]:
    if not DASHSCOPE_API_KEY:
        raise RuntimeError("DASHSCOPE_API_KEY is missing in .env")

    url = f"{WAN_BASE_URL}/tasks/{task_id}"
    headers = {"Authorization": f"Bearer {DASHSCOPE_API_KEY}"}
    response = requests.get(url, headers=headers, timeout=60)
    response.raise_for_status()
    return response.json()


def _story_engine_prompt(payload: Dict[str, Any]) -> str:
    return f"""
你是短剧创作智能体的第一层：故事引擎。
请基于用户输入，输出严格 JSON（不要解释文字）。

用户输入：
- 创意: {payload.get('idea', '')}
- 主题偏好: {payload.get('theme', '')}
- 情绪基调: {payload.get('tone', '')}
- 结构模板偏好: {payload.get('structure', '')}

JSON schema:
{{
  "story_card": {{
    "logline": "一句话故事梗概",
    "theme": "核心主题",
    "tone": "情绪基调",
    "structure_template": "所用结构模板",
    "core_conflict": "核心冲突",
    "anchor_points": ["开端锚点", "转折锚点", "高潮锚点", "结局锚点"],
    "hook": "前三秒抓人钩子",
    "ending_type": "开放式/反转式/治愈式等"
  }},
  "next_questions": ["建议用户补充的问题1", "建议用户补充的问题2"]
}}
""".strip()


def _workshop_prompt(payload: Dict[str, Any]) -> str:
    story_card = payload.get("story_card", {})
    role_requirements = payload.get("role_requirements", "")
    plot_requirements = payload.get("plot_requirements", "")

    return f"""
你是短剧创作智能体的第二层：剧本工坊。
请基于故事卡生成角色、情节节点、对白草稿，输出严格 JSON。

故事卡：
{json.dumps(story_card, ensure_ascii=False, indent=2)}

用户角色要求：{role_requirements}
用户情节要求：{plot_requirements}

JSON schema:
{{
  "characters": [
    {{
      "name": "角色名",
      "tags": ["职业", "性格", "目标", "缺陷"],
      "motivation": "核心动机",
      "arc": "角色弧光"
    }}
  ],
  "relationships": [
    {{"from": "A", "to": "B", "type": "关系类型", "tension": "冲突点"}}
  ],
  "plot_nodes": [
    {{
      "id": "N1",
      "template_stage": "激励事件/第一次转折/高潮等",
      "summary": "节点剧情",
      "consistency_check": "若存在潜在矛盾则提示，否则写无",
      "dialogue_draft": ["角色: 台词"],
      "action_draft": "动作与场面调度"
    }}
  ],
  "timeline_view": ["按时间顺序的节点ID"],
  "card_wall_groups": [
    {{"group": "铺垫/冲突/反转", "node_ids": ["N1", "N2"]}}
  ]
}}
""".strip()


def _storyboard_prompt(payload: Dict[str, Any]) -> str:
    workshop = payload.get("workshop", {})
    style = payload.get("visual_style", "")

    return f"""
你是短剧创作智能体的第三层：分镜工厂。
请将剧本节点转换为分镜卡，输出严格 JSON。

剧本工坊结果：
{json.dumps(workshop, ensure_ascii=False, indent=2)}

视觉风格要求：{style}

JSON schema:
{{
  "storyboards": [
    {{
      "shot_id": "S1",
      "related_node_id": "N1",
      "shot_type": "特写/中景/全景",
      "camera_movement": "固定/推/拉/摇/跟拍",
      "visual_description": "画面内容",
      "dialogue_or_sfx": "对白或音效",
      "duration_sec": 4,
      "shooting_note": "拍摄备注"
    }}
  ],
  "estimated_total_duration_sec": 60,
  "export_ready_checklist": ["服化道", "场景", "收音", "灯光"]
}}
""".strip()


def _command_prompt(payload: Dict[str, Any]) -> str:
    command = payload.get("command", "")
    project_state = payload.get("project_state", {})

    return f"""
你是短剧创作助手的全局指令执行器。
请读取当前状态并执行用户自然语言命令，输出严格 JSON。

用户命令：{command}
当前项目状态：
{json.dumps(project_state, ensure_ascii=False, indent=2)}

JSON schema:
{{
  "command_understanding": "你对命令的理解",
  "updated_state": {{
    "story_card": {{}} ,
    "workshop": {{}} ,
    "storyboard": {{}}
  }},
  "consistency_report": ["一致性检查结果"],
  "suggestions": ["下一步建议"]
}}
""".strip()


def _export_markdown(payload: Dict[str, Any]) -> Dict[str, Any]:
    story_card = payload.get("story_card", {})
    workshop = payload.get("workshop", {})
    storyboard = payload.get("storyboard", {})

    lines: List[str] = []
    lines.append("# AI短剧项目导出")
    lines.append("")
    lines.append("## 1. 故事卡")
    lines.append(f"- Logline: {story_card.get('logline', '')}")
    lines.append(f"- 主题: {story_card.get('theme', '')}")
    lines.append(f"- 基调: {story_card.get('tone', '')}")
    lines.append(f"- 结构模板: {story_card.get('structure_template', '')}")
    lines.append(f"- 核心冲突: {story_card.get('core_conflict', '')}")
    lines.append("")

    lines.append("## 2. 角色设定")
    for c in workshop.get("characters", []):
        tags = ", ".join(c.get("tags", []))
        lines.append(f"- {c.get('name', '未命名角色')} | 标签: {tags} | 动机: {c.get('motivation', '')}")
    lines.append("")

    lines.append("## 3. 情节脉络")
    for n in workshop.get("plot_nodes", []):
        lines.append(
            f"- {n.get('id', '')} [{n.get('template_stage', '')}] {n.get('summary', '')}"
        )
    lines.append("")

    lines.append("## 4. 分镜表")
    lines.append("| 镜头ID | 对应节点 | 景别 | 运镜 | 画面 | 对白/音效 | 时长(秒) |")
    lines.append("|---|---|---|---|---|---|---|")
    for s in storyboard.get("storyboards", []):
        lines.append(
            "| {shot_id} | {node} | {shot_type} | {move} | {visual} | {sound} | {duration} |".format(
                shot_id=s.get("shot_id", ""),
                node=s.get("related_node_id", ""),
                shot_type=s.get("shot_type", ""),
                move=s.get("camera_movement", ""),
                visual=str(s.get("visual_description", "")).replace("|", "\\|"),
                sound=str(s.get("dialogue_or_sfx", "")).replace("|", "\\|"),
                duration=s.get("duration_sec", ""),
            )
        )

    return {"markdown": "\n".join(lines)}


def _build_docx(payload: Dict[str, Any]) -> BytesIO:
    story_card = payload.get("story_card", {})
    workshop = payload.get("workshop", {})
    storyboard = payload.get("storyboard", {})

    doc = Document()
    doc.add_heading("AI短剧项目导出", level=1)

    doc.add_heading("1. 故事卡", level=2)
    doc.add_paragraph(f"Logline: {story_card.get('logline', '')}")
    doc.add_paragraph(f"主题: {story_card.get('theme', '')}")
    doc.add_paragraph(f"基调: {story_card.get('tone', '')}")
    doc.add_paragraph(f"结构模板: {story_card.get('structure_template', '')}")
    doc.add_paragraph(f"核心冲突: {story_card.get('core_conflict', '')}")

    doc.add_heading("2. 角色设定", level=2)
    for c in workshop.get("characters", []):
        tags = ", ".join(c.get("tags", []))
        doc.add_paragraph(
            f"{c.get('name', '未命名角色')} | 标签: {tags} | 动机: {c.get('motivation', '')}",
            style="List Bullet",
        )

    doc.add_heading("3. 情节脉络", level=2)
    for n in workshop.get("plot_nodes", []):
        doc.add_paragraph(
            f"{n.get('id', '')} [{n.get('template_stage', '')}] {n.get('summary', '')}",
            style="List Bullet",
        )

    doc.add_heading("4. 分镜表", level=2)
    table = doc.add_table(rows=1, cols=7)
    headers = ["镜头ID", "对应节点", "景别", "运镜", "画面", "对白/音效", "时长(秒)"]
    header_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        header_cells[i].text = text

    for s in storyboard.get("storyboards", []):
        row = table.add_row().cells
        row[0].text = str(s.get("shot_id", ""))
        row[1].text = str(s.get("related_node_id", ""))
        row[2].text = str(s.get("shot_type", ""))
        row[3].text = str(s.get("camera_movement", ""))
        row[4].text = str(s.get("visual_description", ""))
        row[5].text = str(s.get("dialogue_or_sfx", ""))
        row[6].text = str(s.get("duration_sec", ""))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _register_pdf_font() -> str:
    """Use built-in CJK font to avoid local font file dependency."""
    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        return "STSong-Light"
    except Exception:
        return "Helvetica"


def _draw_wrapped(
    c: canvas.Canvas, text: str, font_name: str, font_size: int, x: float, y: float, width: float
) -> float:
    lines = simpleSplit(str(text), font_name, font_size, width)
    for line in lines:
        c.drawString(x, y, line)
        y -= font_size + 4
    return y


def _build_pdf(payload: Dict[str, Any]) -> BytesIO:
    story_card = payload.get("story_card", {})
    workshop = payload.get("workshop", {})
    storyboard = payload.get("storyboard", {})

    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    font_name = _register_pdf_font()

    def ensure_space(y_pos: float, need: float = 28.0) -> float:
        if y_pos < need:
            c.showPage()
            c.setFont(font_name, 11)
            return height - 50
        return y_pos

    y = height - 50
    c.setFont(font_name, 16)
    c.drawString(40, y, "AI短剧项目导出")
    y -= 34

    c.setFont(font_name, 13)
    c.drawString(40, y, "1. 故事卡")
    y -= 22
    c.setFont(font_name, 11)
    for line in [
        f"Logline: {story_card.get('logline', '')}",
        f"主题: {story_card.get('theme', '')}",
        f"基调: {story_card.get('tone', '')}",
        f"结构模板: {story_card.get('structure_template', '')}",
        f"核心冲突: {story_card.get('core_conflict', '')}",
    ]:
        y = ensure_space(y)
        y = _draw_wrapped(c, line, font_name, 11, 50, y, width - 90)

    y -= 8
    y = ensure_space(y)
    c.setFont(font_name, 13)
    c.drawString(40, y, "2. 角色设定")
    y -= 22
    c.setFont(font_name, 11)
    for ch in workshop.get("characters", []):
        y = ensure_space(y)
        text = (
            f"- {ch.get('name', '未命名角色')} | 标签: {', '.join(ch.get('tags', []))} | "
            f"动机: {ch.get('motivation', '')}"
        )
        y = _draw_wrapped(c, text, font_name, 11, 50, y, width - 90)

    y -= 8
    y = ensure_space(y)
    c.setFont(font_name, 13)
    c.drawString(40, y, "3. 情节脉络")
    y -= 22
    c.setFont(font_name, 11)
    for n in workshop.get("plot_nodes", []):
        y = ensure_space(y)
        text = f"- {n.get('id', '')} [{n.get('template_stage', '')}] {n.get('summary', '')}"
        y = _draw_wrapped(c, text, font_name, 11, 50, y, width - 90)

    y -= 8
    y = ensure_space(y)
    c.setFont(font_name, 13)
    c.drawString(40, y, "4. 分镜表")
    y -= 22
    c.setFont(font_name, 11)
    for s in storyboard.get("storyboards", []):
        y = ensure_space(y, 60)
        c.drawString(50, y, f"{s.get('shot_id', '')} | 节点: {s.get('related_node_id', '')}")
        y -= 16
        y = _draw_wrapped(
            c,
            f"景别: {s.get('shot_type', '')}  运镜: {s.get('camera_movement', '')}",
            font_name,
            11,
            60,
            y,
            width - 110,
        )
        y = _draw_wrapped(
            c,
            f"画面: {s.get('visual_description', '')}",
            font_name,
            11,
            60,
            y,
            width - 110,
        )
        y = _draw_wrapped(
            c,
            f"对白/音效: {s.get('dialogue_or_sfx', '')}  时长: {s.get('duration_sec', '')}秒",
            font_name,
            11,
            60,
            y,
            width - 110,
        )
        y -= 8

    c.save()
    buffer.seek(0)
    return buffer


@app.get("/")
def index() -> str:
    return render_template("index.html")


@app.get("/studio")
def studio() -> str:
    return render_template("studio.html")


@app.get("/visual")
def visual() -> str:
    return render_template("visual.html")


@app.get("/export-center")
def export_center() -> str:
    return render_template("export_center.html")


@app.get("/video-lab")
def video_lab() -> str:
    return render_template("video_lab.html")


@app.post("/api/agent/run")
def run_agent_stage():
    req_json = request.get_json(silent=True) or {}
    stage = req_json.get("stage")
    payload = req_json.get("payload", {})

    if stage not in {"story_engine", "workshop", "storyboard", "command", "export"}:
        return jsonify({"error": "Unsupported stage."}), 400

    try:
        if stage == "story_engine":
            result = _call_qwen_json(
                "你是专业短剧编剧策划，擅长结构化输出。",
                _story_engine_prompt(payload),
            )
        elif stage == "workshop":
            result = _call_qwen_json(
                "你是专业短剧编剧，擅长角色与情节构建，并做一致性检查。",
                _workshop_prompt(payload),
            )
        elif stage == "storyboard":
            result = _call_qwen_json(
                "你是分镜导演，擅长把剧情拆成可拍摄镜头。",
                _storyboard_prompt(payload),
            )
        elif stage == "command":
            result = _call_qwen_json(
                "你是编剧助手，负责执行自然语言编辑命令并保持一致性。",
                _command_prompt(payload),
            )
        else:
            result = _export_markdown(payload)

        return jsonify({"ok": True, "stage": stage, "result": result})
    except requests.HTTPError as e:
        detail: Optional[str] = None
        if e.response is not None:
            detail = e.response.text
        return jsonify({"ok": False, "error": "Qwen API request failed", "detail": detail}), 502
    except Exception as e:  # noqa: BLE001
        return jsonify({"ok": False, "error": str(e)}), 500


@app.post("/api/export/docx")
def export_docx():
    req_json = request.get_json(silent=True) or {}
    payload = req_json.get("payload", req_json)
    try:
        file_obj = _build_docx(payload)
        return send_file(
            file_obj,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="ai_short_drama_export.docx",
        )
    except Exception as e:  # noqa: BLE001
        return jsonify({"ok": False, "error": str(e)}), 500


@app.post("/api/export/pdf")
def export_pdf():
    req_json = request.get_json(silent=True) or {}
    payload = req_json.get("payload", req_json)
    try:
        file_obj = _build_pdf(payload)
        return send_file(
            file_obj,
            mimetype="application/pdf",
            as_attachment=True,
            download_name="ai_short_drama_export.pdf",
        )
    except Exception as e:  # noqa: BLE001
        return jsonify({"ok": False, "error": str(e)}), 500


@app.post("/api/video/script")
def generate_video_script():
    req_json = request.get_json(silent=True) or {}
    payload = req_json.get("payload", req_json)

    try:
        script = _call_qwen_text(
            "你是电影短剧导演和提示词工程师，擅长输出可直接用于视频生成的文本。",
            _video_script_prompt(payload),
        )
        return jsonify({"ok": True, "script": script})
    except requests.HTTPError as e:
        detail: Optional[str] = None
        if e.response is not None:
            detail = e.response.text
        return jsonify({"ok": False, "error": "Qwen API request failed", "detail": detail}), 502
    except Exception as e:  # noqa: BLE001
        return jsonify({"ok": False, "error": str(e)}), 500


@app.post("/api/video/create-task")
def create_video_task():
    req_json = request.get_json(silent=True) or {}
    payload = req_json.get("payload", req_json)
    try:
        result = _create_video_task(payload)
        return jsonify({"ok": True, "result": result})
    except requests.HTTPError as e:
        detail: Optional[str] = None
        if e.response is not None:
            detail = e.response.text
        return jsonify({"ok": False, "error": "Video task creation failed", "detail": detail}), 502
    except Exception as e:  # noqa: BLE001
        return jsonify({"ok": False, "error": str(e)}), 500


@app.get("/api/video/task/<task_id>")
def get_video_task(task_id: str):
    try:
        result = _query_video_task(task_id)
        return jsonify({"ok": True, "result": result})
    except requests.HTTPError as e:
        detail: Optional[str] = None
        if e.response is not None:
            detail = e.response.text
        return jsonify({"ok": False, "error": "Video task query failed", "detail": detail}), 502
    except Exception as e:  # noqa: BLE001
        return jsonify({"ok": False, "error": str(e)}), 500


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8000, debug=True)
